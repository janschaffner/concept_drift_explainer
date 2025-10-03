"""
This module provides functionality for generating and exporting the final drift
analysis results into a formatted, human-readable DOCX report.

It is designed to be called by the frontend, allowing users to download a
professional-looking document that summarizes the key findings for a single
drift. The module includes helpers for sanitizing text to ensure compatibility
with the DOCX format.
"""

import docx
import io
import re
from datetime import datetime

def sanitize_xml_string(text: str) -> str:
    """
    Removes illegal XML characters from a string.

    This is a crucial sanitization step because the underlying `python-docx`
    library works with XML. Certain control characters can create invalid XML,
    causing the document generation to fail. This function removes them to
    ensure the report is always generated successfully.

    Args:
        text: The input string to sanitize.

    Returns:
        The sanitized string.
    """
    if not isinstance(text, str):
        return ""
    # This regex finds and removes characters that are invalid in XML
    illegal_xml_chars_re = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f]')
    return illegal_xml_chars_re.sub('', text)

def generate_docx_report(info: dict, explanation: dict, drift_index: int):
    """
    Generates a formatted DOCX report for a single drift analysis.

    This function assembles a professional report with a clear structure:
    1. Header: Drift type, timeframe, and report title.
    2. Executive Summary: The high-level summary from the Explanation Agent.
    3. Ranked Causes: A detailed breakdown of each potential cause, including
       its confidence score, description, and the supporting evidence snippet.

    The final document is returned as an in-memory byte stream, which is ideal
    for being served as a file download in a web application like Streamlit.

    Args:
        info (dict): The `drift_info` dictionary.
        explanation (dict): The `explanation` dictionary.
        drift_index (int): The 1-based index of the drift for the report title.

    Returns:
        The bytes of the generated DOCX document.
    """
    document = docx.Document()

    # --- Section 1: Report Header ---
    # Sanitize all text before adding it to the document
    dtype = sanitize_xml_string(info.get("drift_type", "Unknown").capitalize())
    start = info.get("start_timestamp", "N/A").split(" ")[0]
    end = info.get("end_timestamp", "N/A").split(" ")[0]
    try:
        s_fmt = datetime.strptime(start, "%Y-%m-%d").strftime("%d %b %Y")
        e_fmt = datetime.strptime(end, "%Y-%m-%d").strftime("%d %b %Y")
        timeframe = f"{s_fmt} – {e_fmt}"
    except Exception:
        timeframe = "N/A"

    document.add_heading(f"Concept Drift Analysis Report: Drift #{drift_index}", level=1)
    document.add_paragraph(f"Drift Type: {dtype}")
    document.add_paragraph(f"Timeframe: {timeframe}")

    # --- Section 2: Executive Summary ---
    summary = sanitize_xml_string(explanation.get("summary", "No summary available."))
    document.add_heading("Executive Summary", level=2)
    document.add_paragraph(summary)
    
    # --- Section 3: Detailed Ranked Causes ---
    document.add_heading("Top Ranked Causes", level=2)
    ranked_causes = explanation.get("ranked_causes", [])
    if not ranked_causes:
        document.add_paragraph("No potential causes were identified.")
    else:
        for i, cause in enumerate(ranked_causes, 1):
            confidence = cause.get('confidence_score', 0) * 100
            # Sanitize all text fields from the 'cause' dictionary
            doc_name = sanitize_xml_string(cause.get('source_document', 'N/A'))
            desc = sanitize_xml_string(cause.get('cause_description', 'N/A'))
            snippet = sanitize_xml_string(cause.get('evidence_snippet', 'No snippet available.'))

            document.add_heading(f"Cause #{i}: {doc_name}", level=3)
            p = document.add_paragraph()
            p.add_run('Confidence: ').bold = True
            p.add_run(f"{confidence:.1f}%")

            p = document.add_paragraph()
            p.add_run('Description: ').bold = True
            p.add_run(desc)

            p = document.add_paragraph()
            p.add_run('Evidence Snippet:').bold = True
            document.add_paragraph(snippet, style='Intense Quote')

    # --- Section 4: Save to an in-memory stream ---
    # Instead of saving to a file on disk, the document is saved to a
    # BytesIO object. This allows to return the raw bytes of the file,
    # which can be used by Streamlit's st.download_button.
    doc_stream = io.BytesIO()
    document.save(doc_stream)
    doc_stream.seek(0)
    
    return doc_stream.getvalue()